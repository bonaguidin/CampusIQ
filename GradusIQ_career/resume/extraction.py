"""Text extraction from uploaded resume files.

Stage 1 of the resume parser: bytes in, plain text out. No AI call, no
database write -- this module exists so that "did we read the file" and "did
the model understand it" stay independently diagnosable.

NEVER RAISES on file content. This is called from a request handler with
untrusted user input, so every failure mode -- encrypted PDF, truncated file,
wrong magic bytes, a .docx that is really a renamed .zip -- comes back as an
ExtractionResult with a non-"ok" status. A raised exception here would 500 the
process on a malformed upload.
"""

import io
from dataclasses import dataclass
from typing import Literal

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


ExtractionStatus = Literal["ok", "empty", "unsupported_format", "extraction_failed"]

PDF_CONTENT_TYPES = frozenset({"application/pdf", "application/x-pdf"})
DOCX_CONTENT_TYPES = frozenset(
    {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
)
LEGACY_DOC_CONTENT_TYPES = frozenset({"application/msword"})

# Legacy .doc is an OLE2 compound file. Detected by magic bytes as well as by
# content type, because browsers routinely mislabel a renamed .doc as .docx --
# without this check that file reaches python-docx and comes back as an opaque
# "extraction_failed" instead of the actionable "re-save as .docx" message.
OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"

# Below this many NON-WHITESPACE characters, the extraction is treated as
# having produced nothing usable. Deliberately not zero: an image-only PDF page
# frequently yields a stray glyph or two, and a scanned resume that extracts to
# "." is the same failure as one that extracts to "" -- both would send an
# effectively blank document to a paid LLM call and get back a confident
# hallucination. No genuine resume is under 20 characters.
MIN_MEANINGFUL_CHARS = 20

# Cell delimiter for DOCX tables. A bare space would make "Acme Corp" and
# "May 2024" read as one run of prose; the pipe keeps the cell boundary legible
# to a downstream prompt.
CELL_DELIMITER = " | "


@dataclass(frozen=True)
class ExtractionResult:
    """Outcome of a single extraction attempt.

    `text` is always a string -- never None -- so callers can log or inspect
    the partial output even on a non-"ok" status.
    """

    text: str
    status: ExtractionStatus
    message: str = ""
    page_count: int | None = None

    @property
    def ok(self) -> bool:
        return self.status == "ok"


def _normalize_content_type(content_type: str) -> str:
    """Lowercase the bare media type, dropping any parameters.

    "Application/PDF; charset=binary" -> "application/pdf"
    """
    if not isinstance(content_type, str):
        return ""
    return content_type.split(";", 1)[0].strip().lower()


def _page_marker(page_number: int) -> str:
    return f"--- page {page_number} ---"


def _meaningful_length(text: str) -> int:
    """Character count ignoring all whitespace."""
    return len("".join(text.split()))


def _extract_pdf(file_bytes: bytes) -> tuple[str, int]:
    """Return (text, page_count) for a PDF, using layout-preserving mode.

    extraction_mode="layout" is not cosmetic. Default mode walks text operators
    in content-stream order, which on a two-column resume interleaves the
    columns line by line -- a left-column employer ends up adjacent to an
    unrelated right-column skill, and the model binds them together. Layout
    mode keeps columns horizontally separated.
    """
    reader = PdfReader(io.BytesIO(file_bytes))
    segments: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text(extraction_mode="layout") or ""
        segments.append(f"{_page_marker(index)}\n{page_text.rstrip()}")
    return "\n\n".join(segments), len(reader.pages)


def _iter_docx_blocks(document: Document):
    """Yield Paragraph and Table objects in true document order.

    document.paragraphs and document.tables are two SEPARATE collections, and
    neither records where the other sat. Concatenating them drops the ordering
    that tells a reader which heading a table belongs under -- and iterating
    only .paragraphs silently omits every table entirely, which on a
    table-formatted resume means the whole work-experience block vanishes with
    no error. Walking the body element is the only way to preserve order.
    """
    body = document.element.body
    paragraph_tag = qn("w:p")
    table_tag = qn("w:tbl")
    for child in body.iterchildren():
        if child.tag == paragraph_tag:
            yield Paragraph(child, document)
        elif child.tag == table_tag:
            yield Table(child, document)


def _render_table(table: Table) -> str:
    lines: list[str] = []
    for row in table.rows:
        # cell.text collapses a cell's own paragraphs with newlines; flatten so
        # one row stays on one line.
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        if any(cells):
            lines.append(CELL_DELIMITER.join(cells))
    return "\n".join(lines)


def _extract_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    segments: list[str] = []
    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                segments.append(text)
        else:
            rendered = _render_table(block)
            if rendered:
                segments.append(rendered)
    return "\n".join(segments)


def extract_resume_text(file_bytes: bytes, content_type: str) -> ExtractionResult:
    """Extract plain text from an uploaded resume.

    Args:
        file_bytes: the raw uploaded bytes.
        content_type: the declared MIME type; parameters (e.g. "; charset=")
            are ignored.

    Returns:
        ExtractionResult. Never raises for bad input -- see module docstring.
    """
    if not file_bytes:
        return ExtractionResult(
            text="",
            status="empty",
            message="The uploaded file is empty (0 bytes).",
        )

    media_type = _normalize_content_type(content_type)

    # Checked before the content-type dispatch so a mislabelled legacy file
    # still gets the actionable message rather than a parser failure.
    if file_bytes.startswith(OLE2_MAGIC) or media_type in LEGACY_DOC_CONTENT_TYPES:
        return ExtractionResult(
            text="",
            status="unsupported_format",
            message=(
                "Legacy Word .doc (Word 97-2003) files are not supported. "
                "Re-save the resume as .docx or PDF and upload it again."
            ),
        )

    if media_type in PDF_CONTENT_TYPES:
        kind = "PDF"
    elif media_type in DOCX_CONTENT_TYPES:
        kind = "DOCX"
    else:
        return ExtractionResult(
            text="",
            status="unsupported_format",
            message=(
                f"Unsupported file type '{media_type or 'unknown'}'. "
                "Upload a PDF or a .docx Word document."
            ),
        )

    page_count: int | None = None
    try:
        if kind == "PDF":
            text, page_count = _extract_pdf(file_bytes)
        else:
            text = _extract_docx(file_bytes)
    except Exception as exc:  # noqa: BLE001 -- untrusted input; see module docstring
        return ExtractionResult(
            text="",
            status="extraction_failed",
            message=f"Could not read the {kind} file: {type(exc).__name__}: {exc}",
        )

    # Measured on the extracted content only. The PDF page markers this module
    # adds are its own output and must not be what pushes a blank scan over the
    # threshold.
    if kind == "PDF":
        body_only = "\n".join(
            line
            for line in text.splitlines()
            if not line.startswith("--- page ")
        )
    else:
        body_only = text

    if _meaningful_length(body_only) < MIN_MEANINGFUL_CHARS:
        return ExtractionResult(
            text=text,
            status="empty",
            message=(
                "No readable text was found in the file. If this is a scanned "
                "or image-only resume, upload a text-based PDF or .docx instead."
            ),
            page_count=page_count,
        )

    return ExtractionResult(text=text, status="ok", page_count=page_count)
