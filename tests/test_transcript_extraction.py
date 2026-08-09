"""Tests for CampusIQ_career.transcript.extraction.

FIXTURE STRATEGY: mirrors tests/test_resume_extraction.py -- every fixture is
GENERATED in-process, not checked in as a binary file, so the geometry and the
table structure under test are readable and diffable in the test source rather
than sealed inside an opaque blob.

The fixtures here are transcript-shaped (course rows: code, title, credits,
grade) because that column structure is the thing the shared extractors have to
survive. reportlab is a dev-group dependency used only to WRITE fixture PDFs;
CampusIQ_career/ never imports it.
"""

import io

import pytest
from docx import Document
from reportlab.lib import pdfencrypt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as rl_canvas

from CampusIQ_career.transcript import (
    TranscriptExtractionResult,
    extract_transcript_text,
)


PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
LEGACY_DOC = "application/msword"

PAGE_W, PAGE_H = letter

# (code, title, credits, grade) -- one semester of a real-shaped transcript.
COURSE_ROWS = [
    ("CSCE 121", "Introduction to Program Design", "4.000", "A"),
    ("MATH 151", "Engineering Mathematics I", "4.000", "B+"),
    ("ENGL 104", "Composition and Rhetoric", "3.000", "A-"),
    ("HIST 105", "History of the United States", "3.000", "B"),
]

# Column x-origins for the tabular PDF fixture. These coordinates ARE the test
# in test_transcript_pdf_keeps_course_rows_intact -- layout mode has to keep the
# four cells of one course on one line instead of interleaving them.
COL_X = (60, 150, 400, 480)


# -- fixture builders ---------------------------------------------------------


def _tabular_transcript_pdf() -> bytes:
    """A course table drawn as positioned text -- the real transcript shape."""
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    y = PAGE_H - 72
    c.drawString(COL_X[0], y, "OFFICIAL ACADEMIC TRANSCRIPT")
    y -= 30
    c.drawString(COL_X[0], y, "Fall 2024")
    y -= 20
    for label, x in zip(("Course", "Title", "Credits", "Grade"), COL_X):
        c.drawString(x, y, label)
    y -= 18
    for code, title, credits, grade in COURSE_ROWS:
        for value, x in zip((code, title, credits, grade), COL_X):
            c.drawString(x, y, value)
        y -= 18
    y -= 12
    c.drawString(COL_X[0], y, "Term GPA: 3.425")
    c.showPage()
    c.save()
    return buf.getvalue()


def _multipage_transcript_pdf() -> bytes:
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, PAGE_H - 72, "Fall 2024 term record: CSCE 121 credits 4.000")
    c.showPage()
    c.drawString(72, PAGE_H - 72, "Spring 2025 term record: CSCE 221 credits 4.000")
    c.showPage()
    c.save()
    return buf.getvalue()


def _scanned_transcript_pdf() -> bytes:
    """Drawn shapes, zero text operators -- a photographed paper transcript."""
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.rect(80, 380, 440, 320, fill=1)
    c.circle(300, 200, 60, fill=1)
    c.showPage()
    c.save()
    return buf.getvalue()


def _near_blank_transcript_pdf() -> bytes:
    """One stray glyph -- what a scan of a faint page actually yields."""
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, PAGE_H - 72, ".")
    c.showPage()
    c.save()
    return buf.getvalue()


def _encrypted_transcript_pdf() -> bytes:
    """Non-empty USER password: genuinely unreadable without the password."""
    enc = pdfencrypt.StandardEncryption("userpw", ownerPassword="ownerpw")
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter, encrypt=enc)
    c.drawString(72, PAGE_H - 72, "CSCE 121 Introduction to Program Design 4.000 A")
    c.showPage()
    c.save()
    return buf.getvalue()


def _permission_restricted_transcript_pdf() -> bytes:
    """Owner password, EMPTY user password -- the registrar's "no printing" PDF.

    reader.is_encrypted is True for this file and the text extracts perfectly.
    It is the reason the module catches FileNotDecryptedError instead of
    checking that flag.
    """
    enc = pdfencrypt.StandardEncryption("", ownerPassword="ownerpw", canPrint=0)
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter, encrypt=enc)
    y = PAGE_H - 72
    for code, title, credits, grade in COURSE_ROWS:
        c.drawString(72, y, f"{code}  {title}  {credits}  {grade}")
        y -= 18
    c.showPage()
    c.save()
    return buf.getvalue()


def _docx_transcript_with_table() -> bytes:
    """Course rows live ONLY in table cells -- the common .docx transcript."""
    d = Document()
    d.add_paragraph("OFFICIAL ACADEMIC TRANSCRIPT")
    d.add_paragraph("Fall 2024")
    table = d.add_table(rows=len(COURSE_ROWS) + 1, cols=4)
    for col, label in enumerate(("Course", "Title", "Credits", "Grade")):
        table.cell(0, col).text = label
    for row_index, row in enumerate(COURSE_ROWS, start=1):
        for col, value in enumerate(row):
            table.cell(row_index, col).text = value
    d.add_paragraph("Term GPA: 3.425")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _docx_paragraphs_around_table() -> bytes:
    d = Document()
    d.add_paragraph("FALL_2024_HEADING")
    table = d.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "CSCE121_IN_CELL"
    table.cell(0, 1).text = "GRADE_A_IN_CELL"
    d.add_paragraph("SPRING_2025_HEADING")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _empty_docx() -> bytes:
    buf = io.BytesIO()
    Document().save(buf)
    return buf.getvalue()


def _fake_legacy_doc() -> bytes:
    """OLE2 compound-file magic -- what a real Word 97-2003 .doc starts with."""
    return b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512


# -- 1. successful PDF extraction ---------------------------------------------


def test_tabular_transcript_pdf_extracts_cleanly():
    result = extract_transcript_text(_tabular_transcript_pdf(), PDF)

    assert isinstance(result, TranscriptExtractionResult)
    assert result.status == "ok"
    assert result.ok is True
    assert result.page_count == 1
    for expected in [
        "OFFICIAL ACADEMIC TRANSCRIPT",
        "Fall 2024",
        "CSCE 121",
        "Introduction to Program Design",
        "Term GPA: 3.425",
    ]:
        assert expected in result.text, f"missing {expected!r}"


def test_transcript_pdf_keeps_course_rows_intact():
    """The layout-mode guarantee, stated in transcript terms.

    A course row is only meaningful as a unit: code, title, credits and grade
    have to stay on one line. Default extraction mode walks text operators in
    stream order and would emit all four course codes, then all four titles,
    then all four grades -- at which point nothing downstream can tell which
    grade belongs to which course.
    """
    result = extract_transcript_text(_tabular_transcript_pdf(), PDF)
    assert result.status == "ok"

    content = [
        line
        for line in result.text.splitlines()
        if line.strip() and not line.startswith("--- page ")
    ]

    for code, title, credits, grade in COURSE_ROWS:
        row = next((line for line in content if code in line), None)
        assert row is not None, f"no line contains {code!r}"
        assert title in row, f"{code}: title split onto another line -- {row!r}"
        assert credits in row, f"{code}: credits split onto another line -- {row!r}"
        assert row.rstrip().endswith(grade), (
            f"{code}: grade {grade!r} is not on its own course's line -- {row!r}"
        )

    # NEGATIVE: the default-mode interleaving must NOT be present.
    assert "CSCE 121\nMATH 151" not in result.text
    assert "A\nB+" not in result.text


def test_pdf_pages_are_separated_by_a_reasonable_marker():
    result = extract_transcript_text(_multipage_transcript_pdf(), PDF)

    assert result.status == "ok"
    assert result.page_count == 2
    assert "--- page 1 ---" in result.text
    assert "--- page 2 ---" in result.text
    assert result.text.index("--- page 1 ---") < result.text.index("Fall 2024")
    assert result.text.index("Fall 2024") < result.text.index("--- page 2 ---")
    assert result.text.index("--- page 2 ---") < result.text.index("Spring 2025")


def test_content_type_parameters_and_case_are_tolerated():
    result = extract_transcript_text(
        _tabular_transcript_pdf(), "Application/PDF; charset=binary"
    )

    assert result.status == "ok"


# -- 2. successful DOCX extraction, including tables --------------------------


def test_docx_transcript_table_cell_content_is_extracted():
    result = extract_transcript_text(_docx_transcript_with_table(), DOCX)
    assert result.status == "ok"

    # Every one of these lives ONLY inside a table cell. The naive
    # "\n".join(p.text for p in doc.paragraphs) approach drops all of them --
    # which on a .docx transcript means every single course disappears.
    for code, title, credits, grade in COURSE_ROWS:
        for cell_only in (code, title, credits):
            assert cell_only in result.text, (
                f"{cell_only!r} lives only in a table cell and was dropped -- "
                "this is the silent whole-course-list-loss bug"
            )

    assert "OFFICIAL ACADEMIC TRANSCRIPT" in result.text
    assert "Term GPA: 3.425" in result.text


def test_docx_transcript_keeps_each_course_row_on_one_line():
    result = extract_transcript_text(_docx_transcript_with_table(), DOCX)

    assert "CSCE 121 | Introduction to Program Design | 4.000 | A" in result.text
    # A bare space delimiter would run the cells together as prose and make
    # "4.000 A" indistinguishable from part of the title.
    assert "CSCE 121Introduction" not in result.text


def test_naive_paragraph_only_extraction_would_have_failed_this():
    """Pins WHY the body walk exists, by demonstrating the alternative."""
    raw = _docx_transcript_with_table()
    naive = "\n".join(p.text for p in Document(io.BytesIO(raw)).paragraphs)

    assert "CSCE 121" not in naive  # the bug, reproduced
    assert "CSCE 121" in extract_transcript_text(raw, DOCX).text  # the fix


def test_docx_preserves_order_of_paragraphs_around_a_table():
    result = extract_transcript_text(_docx_paragraphs_around_table(), DOCX)
    assert result.status == "ok"

    positions = {
        token: result.text.index(token)
        for token in [
            "FALL_2024_HEADING",
            "CSCE121_IN_CELL",
            "GRADE_A_IN_CELL",
            "SPRING_2025_HEADING",
        ]
    }

    # Term headings are paragraphs and course rows are table cells; if the
    # tables were appended after all paragraphs, every course would end up
    # filed under the wrong semester.
    assert positions["FALL_2024_HEADING"] < positions["CSCE121_IN_CELL"]
    assert positions["CSCE121_IN_CELL"] < positions["SPRING_2025_HEADING"]
    assert positions["GRADE_A_IN_CELL"] < positions["SPRING_2025_HEADING"]


# -- 3. scanned / image-only PDF ----------------------------------------------


def test_scanned_pdf_returns_empty_with_the_transcript_specific_message():
    result = extract_transcript_text(_scanned_transcript_pdf(), PDF)

    assert result.status == "empty"
    lowered = result.message.lower()
    assert "scanned" in lowered
    assert "transcript" in lowered
    # Must not promise OCR -- there is none anywhere in this codebase.
    assert "ocr" not in lowered
    # Must not have inherited the resume module's wording.
    assert "resume" not in lowered


def test_near_blank_pdf_returns_empty_not_ok():
    result = extract_transcript_text(_near_blank_transcript_pdf(), PDF)

    assert result.status == "empty"
    assert "scanned" in result.message.lower()


def test_scanned_pdf_page_markers_do_not_defeat_the_threshold():
    """Page markers are this module's own output, not extracted content.

    A 3-page scan carries ~42 characters of markers -- more than
    MIN_MEANINGFUL_CHARS. Counting them would report "ok" on a document with
    zero readable courses in it.
    """
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    for _ in range(3):
        c.rect(100, 400, 200, 100, fill=1)
        c.showPage()
    c.save()

    result = extract_transcript_text(buf.getvalue(), PDF)

    assert result.page_count == 3
    assert result.status == "empty"


# -- 4. encrypted / password-protected PDF ------------------------------------


def test_encrypted_pdf_gets_its_own_status_and_actionable_message():
    result = extract_transcript_text(_encrypted_transcript_pdf(), PDF)

    assert result.status == "encrypted", (
        "a password-protected transcript must be distinguishable from a "
        "corrupt one -- it is the one failure the student can fix themselves"
    )
    lowered = result.message.lower()
    assert "password-protected" in lowered
    assert "unprotected" in lowered
    assert result.text == ""
    # Specifically NOT the generic dump of an exception type.
    assert "FileNotDecryptedError" not in result.message


def test_permission_restricted_pdf_still_extracts_normally():
    """The is_encrypted trap.

    This file reports reader.is_encrypted == True but has an EMPTY user
    password, so its text is readable. Registrars issue exactly this (owner
    password set to block printing/copying). A pre-check on is_encrypted would
    reject a perfectly good transcript.
    """
    result = extract_transcript_text(_permission_restricted_transcript_pdf(), PDF)

    assert result.status == "ok"
    assert "CSCE 121" in result.text
    assert "Introduction to Program Design" in result.text


def test_encrypted_is_not_conflated_with_corrupt():
    encrypted = extract_transcript_text(_encrypted_transcript_pdf(), PDF)
    corrupt = extract_transcript_text(b"%PDF-1.7\n" + b"\x00\xff\xde\xad" * 50, PDF)

    assert encrypted.status == "encrypted"
    assert corrupt.status == "extraction_failed"
    assert encrypted.message != corrupt.message


# -- 5. unsupported formats ---------------------------------------------------


@pytest.mark.parametrize(
    "content_type, payload",
    [
        (LEGACY_DOC, _fake_legacy_doc()),
        (LEGACY_DOC, b"anything at all"),
        # Mislabelled: a real .doc sent with the .docx content type. The magic
        # bytes still identify it, so the message stays actionable.
        (DOCX, _fake_legacy_doc()),
        (PDF, _fake_legacy_doc()),
    ],
)
def test_legacy_doc_is_unsupported_not_a_crash(content_type, payload):
    result = extract_transcript_text(payload, content_type)

    assert result.status == "unsupported_format"
    assert ".doc" in result.message
    assert "docx" in result.message.lower()
    assert "transcript" in result.message.lower()
    assert result.text == ""


@pytest.mark.parametrize(
    "content_type",
    ["text/plain", "image/png", "application/zip", "application/octet-stream", ""],
)
def test_other_unknown_content_types_are_unsupported(content_type):
    result = extract_transcript_text(b"some bytes here", content_type)

    assert result.status == "unsupported_format"
    assert "PDF" in result.message
    assert "transcript" in result.message.lower()


# -- 6. corrupt input never raises --------------------------------------------


@pytest.mark.parametrize(
    "label, payload",
    [
        (
            "truncated-half",
            _tabular_transcript_pdf()[: len(_tabular_transcript_pdf()) // 2],
        ),
        ("truncated-tiny", _tabular_transcript_pdf()[:40]),
        ("magic-then-garbage", b"%PDF-1.7\n" + b"\x00\xff\xde\xad" * 50),
        ("not-a-pdf-at-all", b"this is plain text, definitely not a pdf"),
    ],
)
def test_corrupt_pdf_returns_extraction_failed_without_raising(label, payload):
    result = extract_transcript_text(payload, PDF)

    assert result.status == "extraction_failed", label
    assert result.message
    assert result.text == ""


@pytest.mark.parametrize(
    "label, payload",
    [
        ("truncated-docx", _docx_transcript_with_table()[:100]),
        ("zip-but-not-docx", b"PK\x03\x04" + b"\x00" * 100),
        ("random-bytes", bytes(range(256))),
    ],
)
def test_corrupt_docx_returns_extraction_failed_without_raising(label, payload):
    result = extract_transcript_text(payload, DOCX)

    assert result.status == "extraction_failed", label
    assert result.message
    assert result.text == ""


def test_extraction_never_raises_across_a_sweep_of_hostile_inputs():
    """Nothing reaching a request handler may escape as an exception."""
    payloads = [
        b"\x00",
        b"\xff" * 1000,
        b"%PDF",
        b"PK",
        bytes(range(256)) * 4,
        "trànscript non-ascii text".encode(),
        b"<html><body>not a transcript</body></html>",
    ]
    types = [PDF, DOCX, LEGACY_DOC, "text/plain", "", "application/pdf; x=1"]

    for payload in payloads:
        for content_type in types:
            result = extract_transcript_text(payload, content_type)
            assert result.status in {
                "ok",
                "empty",
                "unsupported_format",
                "extraction_failed",
                "encrypted",
            }
            assert isinstance(result.text, str)


# -- 7. empty input -----------------------------------------------------------


@pytest.mark.parametrize("content_type", [PDF, DOCX, LEGACY_DOC, "text/plain", ""])
def test_empty_bytes_returns_empty_for_every_content_type(content_type):
    result = extract_transcript_text(b"", content_type)

    assert result.status == "empty"
    assert result.text == ""
    assert "empty" in result.message.lower()


def test_docx_with_no_content_returns_empty():
    result = extract_transcript_text(_empty_docx(), DOCX)

    assert result.status == "empty"


# -- 8. no truncation happens in stage 1 --------------------------------------


def test_stage_one_does_not_truncate_long_transcripts():
    """Stage 1 hands Stage 2 the WHOLE document.

    A four-year transcript comfortably exceeds the resume parser's 60k prompt
    cap. Truncation is Stage 2's decision to make -- and per the audit it must
    be a hard error there, because a silently shortened transcript drops
    courses and yields a GPA computed over a subset of them. Stage 1 must not
    pre-empt that by dropping anything itself.
    """
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    for page in range(12):
        y = PAGE_H - 72
        for row in range(35):
            c.drawString(60, y, f"CSCE {page:02d}{row:02d}  Course Title Filler  3.000  A")
            y -= 18
        c.showPage()
    c.save()

    result = extract_transcript_text(buf.getvalue(), PDF)

    assert result.status == "ok"
    assert result.page_count == 12
    # First and LAST course both survive -- nothing was cut off the end.
    assert "CSCE 0000" in result.text
    assert "CSCE 1134" in result.text
    assert "truncated" not in result.text.lower()
