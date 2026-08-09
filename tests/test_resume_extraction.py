"""Tests for GradusIQ_career.resume.extraction.

FIXTURE STRATEGY: every fixture is GENERATED in-process, not checked in as a
binary file. Two reasons, both specific to what these tests assert:

  1. The two-column PDF test turns entirely on column geometry -- the x
     coordinates 60 and 340 are the thing under test. In a checked-in PDF that
     geometry is invisible to a reviewer and undiffable in git; written as
     reportlab calls it is right there in the test.
  2. The DOCX table tests assert that content living ONLY in a table cell
     survives. Building the table with python-docx makes "only in a cell"
     literal and verifiable, rather than a claim about an opaque blob.

reportlab is a dev-group dependency for exactly this purpose and is never
imported by GradusIQ_career/.
"""

import io

import pytest
from docx import Document
from reportlab.lib import pdfencrypt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as rl_canvas

from GradusIQ_career.resume import ExtractionResult, extract_resume_text


PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
LEGACY_DOC = "application/msword"

PAGE_W, PAGE_H = letter


# ── fixture builders ─────────────────────────────────────────────────────────


def _single_column_pdf() -> bytes:
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    y = PAGE_H - 72
    for line in [
        "Jane Doe",
        "jane.doe@example.edu",
        "EXPERIENCE",
        "Acme Corp - Software Intern",
        "May 2024 - August 2024",
        "Built an internal analytics dashboard.",
        "EDUCATION",
        "Texas A&M University, B.S. Computer Science",
    ]:
        c.drawString(72, y, line)
        y -= 18
    c.showPage()
    c.save()
    return buf.getvalue()


def _two_column_pdf() -> bytes:
    """Left and right columns sharing baselines -- the interleaving trap.

    In pypdf's DEFAULT extraction mode this produces
    "EXPERIENCE\\nSKILLS\\nAcme Corp\\nPython..." because it walks text
    operators in stream order. Layout mode keeps the two columns side by side
    on one line. That difference is what test 2 pins.
    """
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    left = [
        "EXPERIENCE",
        "Acme Corp",
        "Software Intern",
        "May 2024 - Aug 2024",
        "Built a dashboard.",
    ]
    right = ["SKILLS", "Python", "SQL", "React", "Docker"]
    y = PAGE_H - 72
    for i in range(max(len(left), len(right))):
        if i < len(left):
            c.drawString(60, y, left[i])
        if i < len(right):
            c.drawString(340, y, right[i])
        y -= 18
    c.showPage()
    c.save()
    return buf.getvalue()


def _multipage_pdf() -> bytes:
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, PAGE_H - 72, "Page one content: Jane Doe resume header")
    c.showPage()
    c.drawString(72, PAGE_H - 72, "Page two content: additional project history")
    c.showPage()
    c.save()
    return buf.getvalue()


def _image_only_pdf() -> bytes:
    """A page with drawn shapes and zero text operators -- a scanned resume."""
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.rect(100, 400, 200, 100, fill=1)
    c.circle(300, 200, 50, fill=1)
    c.showPage()
    c.save()
    return buf.getvalue()


def _near_blank_pdf() -> bytes:
    """One stray character -- functionally the same failure as a blank page."""
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, PAGE_H - 72, ".")
    c.showPage()
    c.save()
    return buf.getvalue()


def _encrypted_pdf() -> bytes:
    enc = pdfencrypt.StandardEncryption("userpw", ownerPassword="ownerpw")
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter, encrypt=enc)
    c.drawString(72, PAGE_H - 72, "secret resume text that cannot be read")
    c.showPage()
    c.save()
    return buf.getvalue()


def _docx_with_table() -> bytes:
    """Work history lives ONLY in table cells -- never in a paragraph."""
    d = Document()
    d.add_paragraph("Jane Doe")
    d.add_paragraph("EXPERIENCE")
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Acme Corp"
    table.cell(0, 1).text = "May 2024 - Aug 2024"
    table.cell(1, 0).text = "Software Intern"
    table.cell(1, 1).text = "Dallas, TX"
    d.add_paragraph("SKILLS")
    d.add_paragraph("Python, SQL, React")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _docx_paragraphs_around_table() -> bytes:
    d = Document()
    d.add_paragraph("ALPHA_BEFORE")
    table = d.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "BRAVO_IN_CELL"
    table.cell(0, 1).text = "CHARLIE_IN_CELL"
    d.add_paragraph("DELTA_AFTER")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _empty_docx() -> bytes:
    buf = io.BytesIO()
    Document().save(buf)
    return buf.getvalue()


def _fake_legacy_doc() -> bytes:
    """OLE2 compound-file magic -- what a real Word 97-2003 .doc starts with."""
    return b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512


# ── 1. single-column PDF ─────────────────────────────────────────────────────


def test_single_column_pdf_extracts_cleanly():
    result = extract_resume_text(_single_column_pdf(), PDF)

    assert isinstance(result, ExtractionResult)
    assert result.status == "ok"
    assert result.ok is True
    assert result.page_count == 1
    for expected in [
        "Jane Doe",
        "jane.doe@example.edu",
        "EXPERIENCE",
        "Acme Corp",
        "Software Intern",
        "Texas A&M University",
    ]:
        assert expected in result.text, f"missing {expected!r}"


def test_pdf_pages_are_separated_by_a_reasonable_marker():
    result = extract_resume_text(_multipage_pdf(), PDF)

    assert result.status == "ok"
    assert result.page_count == 2
    assert "--- page 1 ---" in result.text
    assert "--- page 2 ---" in result.text
    # The marker must precede its own page's content, in order.
    assert result.text.index("--- page 1 ---") < result.text.index("Page one content")
    assert result.text.index("Page one content") < result.text.index("--- page 2 ---")
    assert result.text.index("--- page 2 ---") < result.text.index("Page two content")


# ── 2. two-column PDF uses layout mode ───────────────────────────────────────


def test_two_column_pdf_uses_layout_mode_not_default_interleaving():
    result = extract_resume_text(_two_column_pdf(), PDF)
    assert result.status == "ok"

    lines = [line for line in result.text.splitlines() if line.strip()]
    content = [line for line in lines if not line.startswith("--- page ")]

    # LAYOUT MODE SIGNATURE: the two column headings share one physical line,
    # horizontally separated. Default mode would put them on separate lines.
    header = next((line for line in content if "EXPERIENCE" in line), None)
    assert header is not None
    assert "SKILLS" in header, (
        "EXPERIENCE and SKILLS must appear on the same line (layout mode). "
        f"Got lines: {content!r}"
    )
    assert header.index("EXPERIENCE") < header.index("SKILLS")
    # Separated by real horizontal whitespace, not concatenated.
    between = header[header.index("EXPERIENCE") + len("EXPERIENCE"): header.index("SKILLS")]
    assert between.strip() == ""
    assert len(between) > 1

    # Each row keeps its own two cells together on one line.
    row = next((line for line in content if "Acme Corp" in line), None)
    assert row is not None and "Python" in row

    # NEGATIVE: the default-mode interleaving must NOT be present.
    assert "EXPERIENCE\nSKILLS" not in result.text
    assert "Acme Corp\nPython" not in result.text


# ── 3. image-only PDF ────────────────────────────────────────────────────────


def test_image_only_pdf_returns_empty():
    result = extract_resume_text(_image_only_pdf(), PDF)

    assert result.status == "empty"
    assert "no readable text" in result.message.lower()


def test_near_blank_pdf_returns_empty_not_ok():
    """A single stray character is the same failure as a blank page."""
    result = extract_resume_text(_near_blank_pdf(), PDF)

    assert result.status == "empty"


# ── 4. DOCX table content survives ───────────────────────────────────────────


def test_docx_table_cell_content_is_extracted():
    result = extract_resume_text(_docx_with_table(), DOCX)
    assert result.status == "ok"

    # These strings exist ONLY inside table cells. The naive
    # "\n".join(p.text for p in doc.paragraphs) approach drops all four.
    for cell_only in ["Acme Corp", "May 2024 - Aug 2024", "Software Intern", "Dallas, TX"]:
        assert cell_only in result.text, (
            f"{cell_only!r} lives only in a table cell and was dropped -- "
            "this is the silent work-experience-loss bug"
        )

    # Paragraph content is still present too.
    assert "Jane Doe" in result.text
    assert "SKILLS" in result.text


def test_docx_table_cells_use_a_non_prose_delimiter():
    result = extract_resume_text(_docx_with_table(), DOCX)

    assert "Acme Corp | May 2024 - Aug 2024" in result.text
    # A bare space would run the two cells together as prose.
    assert "Acme CorpMay 2024" not in result.text
    assert "Acme Corp May 2024" not in result.text


def test_naive_paragraph_only_extraction_would_have_failed_this():
    """Pins WHY the body walk exists, by demonstrating the alternative."""
    raw = _docx_with_table()
    naive = "\n".join(p.text for p in Document(io.BytesIO(raw)).paragraphs)

    assert "Acme Corp" not in naive  # the bug, reproduced
    assert "Acme Corp" in extract_resume_text(raw, DOCX).text  # the fix


# ── 5. DOCX document order ───────────────────────────────────────────────────


def test_docx_preserves_order_of_paragraphs_around_a_table():
    result = extract_resume_text(_docx_paragraphs_around_table(), DOCX)
    assert result.status == "ok"

    positions = {
        token: result.text.index(token)
        for token in ["ALPHA_BEFORE", "BRAVO_IN_CELL", "CHARLIE_IN_CELL", "DELTA_AFTER"]
    }

    assert positions["ALPHA_BEFORE"] < positions["BRAVO_IN_CELL"], (
        "the paragraph before the table must come first"
    )
    assert positions["BRAVO_IN_CELL"] < positions["DELTA_AFTER"], (
        "table content must sit between the paragraphs, not be appended after them"
    )
    assert positions["CHARLIE_IN_CELL"] < positions["DELTA_AFTER"]


# ── 6. legacy .doc ───────────────────────────────────────────────────────────


@pytest.mark.parametrize(
    "content_type, payload",
    [
        (LEGACY_DOC, _fake_legacy_doc()),
        (LEGACY_DOC, b"anything at all"),
        # Mislabelled: a real .doc sent with the .docx content type. The magic
        # bytes still identify it, so the message stays actionable.
        (DOCX, _fake_legacy_doc()),
        (PDF, _fake_legacy_doc()),
    ],
)
def test_legacy_doc_is_unsupported_not_a_crash(content_type, payload):
    result = extract_resume_text(payload, content_type)

    assert result.status == "unsupported_format"
    assert ".doc" in result.message
    assert "docx" in result.message.lower()
    assert result.text == ""


@pytest.mark.parametrize(
    "content_type",
    ["text/plain", "image/png", "application/zip", "application/octet-stream", ""],
)
def test_other_unknown_content_types_are_unsupported(content_type):
    result = extract_resume_text(b"some bytes here", content_type)

    assert result.status == "unsupported_format"
    assert "PDF" in result.message


def test_content_type_parameters_and_case_are_tolerated():
    result = extract_resume_text(_single_column_pdf(), "Application/PDF; charset=binary")

    assert result.status == "ok"


# ── 7. corrupt input never raises ────────────────────────────────────────────


@pytest.mark.parametrize(
    "label, payload",
    [
        ("truncated-half", _single_column_pdf()[: len(_single_column_pdf()) // 2]),
        ("truncated-tiny", _single_column_pdf()[:40]),
        ("magic-then-garbage", b"%PDF-1.7\n" + b"\x00\xff\xde\xad" * 50),
        ("not-a-pdf-at-all", b"this is plain text, definitely not a pdf"),
        ("encrypted", _encrypted_pdf()),
    ],
)
def test_corrupt_pdf_returns_extraction_failed_without_raising(label, payload):
    result = extract_resume_text(payload, PDF)

    assert result.status == "extraction_failed", label
    assert result.message
    assert result.text == ""


@pytest.mark.parametrize(
    "label, payload",
    [
        ("truncated-docx", _docx_with_table()[:100]),
        ("zip-but-not-docx", b"PK\x03\x04" + b"\x00" * 100),
        ("random-bytes", bytes(range(256))),
    ],
)
def test_corrupt_docx_returns_extraction_failed_without_raising(label, payload):
    result = extract_resume_text(payload, DOCX)

    assert result.status == "extraction_failed", label
    assert result.message
    assert result.text == ""


def test_extraction_never_raises_across_a_sweep_of_hostile_inputs():
    """Nothing reaching a request handler may escape as an exception."""
    payloads = [
        b"\x00",
        b"\xff" * 1000,
        b"%PDF",
        b"PK",
        bytes(range(256)) * 4,
        "résumé non-ascii text".encode(),
        b"<html><body>not a resume</body></html>",
    ]
    types = [PDF, DOCX, LEGACY_DOC, "text/plain", "", "application/pdf; x=1"]

    for payload in payloads:
        for content_type in types:
            result = extract_resume_text(payload, content_type)
            assert result.status in {
                "ok",
                "empty",
                "unsupported_format",
                "extraction_failed",
            }
            assert isinstance(result.text, str)


# ── 8. empty input ───────────────────────────────────────────────────────────


@pytest.mark.parametrize("content_type", [PDF, DOCX, LEGACY_DOC, "text/plain", ""])
def test_empty_bytes_returns_empty_for_every_content_type(content_type):
    result = extract_resume_text(b"", content_type)

    assert result.status == "empty"
    assert result.text == ""
    assert "empty" in result.message.lower()


def test_docx_with_no_content_returns_empty():
    result = extract_resume_text(_empty_docx(), DOCX)

    assert result.status == "empty"
