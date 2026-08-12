"""Tests for POST /api/v2/student/me/resume/upload and .../career/confirm.

No network: the OpenRouter client is faked, and Supabase is a small in-memory
double that enforces the three things these routes actually depend on --
per-student row visibility (RLS), the natural-key unique indexes, and the
"unconfirmed only" filter. Faking those behaviors rather than stubbing the
calls is what lets a duplicate-skip or a cross-student leak actually fail here.

The natural keys and their NULLS NOT DISTINCT semantics were verified against
the live database during the Stage 1 audit; FakeSupabase reproduces them.
"""

import io
import json

import pytest
from docx import Document
from fastapi.testclient import TestClient
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas as rl_canvas

from GradusIQ_career import api
from GradusIQ_career.ai.types import AIResponse
from GradusIQ_career.profile_builder import build_student_intelligence_profile


TEST_PROXY_SECRET = "test-proxy-secret"
PROXY_HEADERS = {api.PROXY_SECRET_HEADER: TEST_PROXY_SECRET}
AUTH = {"Authorization": "Bearer real-session-jwt"}
HEADERS = {**PROXY_HEADERS, **AUTH}

STUDENT_A = "8f14e45f-ceea-467a-9f0e-1c2d3e4f5a6b"
STUDENT_B = "1a2b3c4d-5e6f-4a7b-8c9d-0e1f2a3b4c5d"

PDF = "application/pdf"
DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

UPLOAD = "/api/v2/student/me/resume/upload"
CONFIRM = "/api/v2/student/me/career/confirm"

NATURAL_KEYS = {
    "certifications": ("student_id", "name"),
    "work_experience": ("student_id", "employer", "role"),
    "projects": ("student_id", "name"),
}


# ── fixtures: real files, generated in-test (mirrors Stage 1) ────────────────


def make_pdf(lines=None) -> bytes:
    lines = lines or [
        "Jane Doe",
        "OBJECTIVE: Software Engineering Intern",
        "EXPERIENCE",
        "Acme Corp - Software Intern - May 2024 to Aug 2024",
        "Built an internal analytics dashboard using Python.",
        "CERTIFICATIONS",
        "AWS Cloud Practitioner, Amazon, 2024",
        "PROJECTS",
        "Campus Scheduler - a React app for course planning",
    ]
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    y = letter[1] - 72
    for line in lines:
        c.drawString(72, y, line)
        y -= 18
    c.showPage()
    c.save()
    return buf.getvalue()


def make_docx() -> bytes:
    d = Document()
    d.add_paragraph("Jane Doe")
    d.add_paragraph("EXPERIENCE")
    t = d.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "Acme Corp"
    t.cell(0, 1).text = "May 2024 - Aug 2024"
    d.add_paragraph("Skills: Python, SQL")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def image_only_pdf() -> bytes:
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    c.rect(100, 400, 200, 100, fill=1)
    c.showPage()
    c.save()
    return buf.getvalue()


# ── fake AI ──────────────────────────────────────────────────────────────────

GOOD_PARSE = {
    "status": "ok",
    "academics": {
        "major_current": "Computer Engineering",
        "expected_graduation": "May 2029",
    },
    "profile": {
        "target_roles": ["Software Engineering Intern"],
        "interests": ["backend"],
        "skills_technical": ["Python", "SQL"],
        "skills_soft": ["communication"],
    },
    "certifications": [
        {"name": "AWS Cloud Practitioner", "issuer": "Amazon", "status": "completed", "date": "2024"}
    ],
    "work_experience": [
        {
            "employer": "Acme Corp",
            "role": "Software Intern",
            "duration": "May 2024 - Aug 2024",
            "location": None,
            "description": "Built an internal analytics dashboard.",
            "skills_gained": ["Python"],
        }
    ],
    "projects": [
        {"name": "Campus Scheduler", "timeframe": None, "description": "React app", "tools": ["React"]}
    ],
}


class FakeAI:
    """Records every call so 'the model was never invoked' is assertable."""

    def __init__(self, text=None):
        self.text = json.dumps(GOOD_PARSE) if text is None else text
        self.calls = []

    def complete(self, **kwargs):
        self.calls.append(kwargs)
        return AIResponse(text=self.text, raw={"choices": []}, model="fake-parsing-model")


class ExplodingAI:
    def complete(self, **kwargs):
        raise AssertionError("the AI client must not be called on this path")


# ── fake Supabase ────────────────────────────────────────────────────────────


class FakeDB:
    """Shared row store. One instance can back several students' clients."""

    def __init__(self):
        self.tables = {
            "students": [],
            "career_profiles": [],
            "certifications": [],
            "work_experience": [],
            "projects": [],
            "student_institutions": [],
            "institutions": [],
            "academic_terms": [],
            "course_records": [],
        }
        self._next = 0

    def new_id(self, prefix):
        self._next += 1
        return f"{prefix}-{self._next:04d}"

    def add_student(self, student_id):
        self.tables["students"].append({"id": student_id, "name": f"Student {student_id[:4]}"})

    def rows_for(self, table, student_id):
        return [r for r in self.tables[table] if r.get("student_id") == student_id]


class FakeQuery:
    def __init__(self, db, table, student_id):
        self.db = db
        self.table_name = table
        self.student_id = student_id
        self.op = None
        self.payload = None
        self.filters = []
        self.on_conflict = None
        self.ignore_duplicates = False

    # -- verbs
    def select(self, *a, **k):
        self.op = "select"
        return self

    def insert(self, json, **k):
        self.op = "insert"
        self.payload = json
        return self

    def upsert(self, json, *, ignore_duplicates=False, on_conflict="", **k):
        self.op = "upsert"
        self.payload = json
        self.ignore_duplicates = ignore_duplicates
        self.on_conflict = on_conflict
        return self

    def update(self, json, **k):
        self.op = "update"
        self.payload = json
        return self

    # -- filters
    def eq(self, column, value):
        self.filters.append(("eq", column, value))
        return self

    def is_(self, column, value):
        self.filters.append(("is", column, value))
        return self

    def in_(self, column, values):
        self.filters.append(("in", column, list(values)))
        return self

    def _matches(self, row):
        for kind, column, value in self.filters:
            if kind == "eq" and row.get(column) != value:
                return False
            if kind == "is":
                wanted = None if value in (None, "null") else value
                if row.get(column) is not wanted:
                    return False
            if kind == "in" and row.get(column) not in value:
                return False
        return True

    def _visible(self):
        """RLS: a session sees only its own student's rows."""
        rows = self.db.tables[self.table_name]
        if self.table_name == "students":
            return [r for r in rows if r["id"] == self.student_id]
        return [r for r in rows if r.get("student_id") == self.student_id]

    def execute(self):
        if self.op == "select":
            return _Result([dict(r) for r in self._visible() if self._matches(r)])

        if self.op == "insert":
            row = dict(self.payload)
            row.setdefault("id", self.db.new_id(self.table_name[:4]))
            self.db.tables[self.table_name].append(row)
            return _Result([dict(row)])

        if self.op == "upsert":
            columns = tuple(c.strip() for c in (self.on_conflict or "").split(",") if c.strip())
            expected = NATURAL_KEYS.get(self.table_name)
            assert columns == expected, (
                f"{self.table_name}: on_conflict {columns} does not match the live "
                f"natural key {expected} -- Postgres would raise 42P10"
            )
            row = dict(self.payload)
            # NULLS NOT DISTINCT: None compares equal to None, matching the
            # live work_experience index.
            key = tuple(row.get(c) for c in columns)
            for existing in self.db.tables[self.table_name]:
                if tuple(existing.get(c) for c in columns) == key:
                    # ON CONFLICT DO NOTHING -> no row returned, none modified.
                    assert self.ignore_duplicates, "expected ignore_duplicates=True"
                    return _Result([])
            row.setdefault("id", self.db.new_id(self.table_name[:4]))
            self.db.tables[self.table_name].append(row)
            return _Result([dict(row)])

        if self.op == "update":
            updated = []
            for row in self._visible():
                if self._matches(row):
                    row.update(self.payload)
                    updated.append(dict(row))
            return _Result(updated)

        raise AssertionError(f"unsupported op {self.op}")


class _Result:
    def __init__(self, data):
        self.data = data


class _Postgrest:
    def __init__(self):
        self.tokens = []

    def auth(self, token):
        self.tokens.append(token)


class FakeSupabase:
    def __init__(self, db, student_id):
        self.db = db
        self.student_id = student_id
        self.postgrest = _Postgrest()

    def table(self, name):
        return FakeQuery(self.db, name, self.student_id)


# ── wiring ───────────────────────────────────────────────────────────────────


def make_test_config(**overrides):
    values = {
        "proxy_secret": TEST_PROXY_SECRET,
        "allowed_origins": ("https://frontend.example",),
        "rate_limit_requests": 100,
        "rate_limit_window_seconds": 60.0,
        "max_concurrent_ai_requests": 2,
    }
    values.update(overrides)
    return api.APIConfig(**values)


@pytest.fixture
def db():
    store = FakeDB()
    store.add_student(STUDENT_A)
    store.add_student(STUDENT_B)
    return store


@pytest.fixture
def client():
    return TestClient(api.create_app(make_test_config()), headers=PROXY_HEADERS)


def patch_session(monkeypatch, db, student_id=STUDENT_A):
    monkeypatch.setattr(api, "build_client_for_token", lambda token: FakeSupabase(db, student_id))


def upload(client, content=None, filename="resume.pdf", content_type=PDF, headers=None):
    return client.post(
        UPLOAD,
        headers=HEADERS if headers is None else headers,
        files={"file": (filename, make_pdf() if content is None else content, content_type)},
    )


# ── 1. happy path: rows land with correct source / confirmed_at ─────────────


def test_pdf_upload_writes_rows_with_resume_source_and_unconfirmed(client, db, monkeypatch):
    patch_session(monkeypatch, db)
    fake = FakeAI()
    monkeypatch.setattr(api, "build_client", lambda: fake)

    response = upload(client)

    assert response.status_code == 200, response.text
    body = response.json()
    assert body["status"] == "ok"
    assert body["extraction"]["status"] == "ok"
    assert body["career_profile"]["outcome"] == "created"
    assert body["written"] == {
        "certifications": {"inserted": 1, "skipped_duplicate": 0},
        "work_experience": {"inserted": 1, "skipped_duplicate": 0},
        "projects": {"inserted": 1, "skipped_duplicate": 0},
    }

    # The model was called once, on the parsing role, with the extracted text.
    assert len(fake.calls) == 1
    assert fake.calls[0]["role"] == "parsing"
    assert "Acme Corp" in fake.calls[0]["messages"][1]["content"]

    profile = db.tables["career_profiles"][0]
    assert profile["student_id"] == STUDENT_A
    assert profile["source"] == "resume_parse"
    assert profile["confirmed_at"] is None
    assert profile["target_roles"] == ["Software Engineering Intern"]
    assert profile["skills_technical"] == ["Python", "SQL"]

    for table in ("certifications", "work_experience", "projects"):
        rows = db.rows_for(table, STUDENT_A)
        assert len(rows) == 1, table
        assert rows[0]["source"] == "resume_parse", table
        assert rows[0]["confirmed_at"] is None, table
        assert rows[0]["career_profile_id"] == profile["id"], table

    cert = db.rows_for("certifications", STUDENT_A)[0]
    assert cert["name"] == "AWS Cloud Practitioner"
    assert cert["status"] == "completed"
    job = db.rows_for("work_experience", STUDENT_A)[0]
    assert job["employer"] == "Acme Corp"
    assert job["skills_gained"] == ["Python"]


def test_docx_upload_also_writes_rows(client, db, monkeypatch):
    patch_session(monkeypatch, db)
    monkeypatch.setattr(api, "build_client", lambda: FakeAI())

    response = upload(client, content=make_docx(), filename="resume.docx", content_type=DOCX)

    assert response.status_code == 200, response.text
    assert response.json()["status"] == "ok"
    assert len(db.rows_for("work_experience", STUDENT_A)) == 1


# ── 2. extraction failure short-circuits before any AI call ─────────────────


@pytest.mark.parametrize(
    ("label", "content", "filename", "content_type", "expected_status"),
    [
        ("image-only pdf", image_only_pdf(), "scan.pdf", PDF, 422),
        ("empty file", b"", "empty.pdf", PDF, 422),
        ("legacy doc", b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 64, "old.doc",
         "application/msword", 415),
        ("unsupported type", b"hello there", "notes.txt", "text/plain", 415),
        ("corrupt pdf", b"%PDF-1.7\n" + b"\xde\xad\xbe\xef" * 20, "broken.pdf", PDF, 422),
    ],
)
def test_extraction_failure_never_reaches_the_model(
    label, content, filename, content_type, expected_status, client, db, monkeypatch
):
    patch_session(monkeypatch, db)
    # Any call at all fails the test.
    monkeypatch.setattr(api, "build_client", lambda: ExplodingAI())

    response = upload(client, content=content, filename=filename, content_type=content_type)

    assert response.status_code == expected_status, f"{label}: {response.text}"
    detail = response.json()["detail"]
    assert detail["error"] == "extraction_failed"
    assert detail["extraction_status"] in {"empty", "unsupported_format", "extraction_failed"}
    assert detail["message"]
    # Nothing was written.
    for table in ("career_profiles", "certifications", "work_experience", "projects"):
        assert db.tables[table] == [], f"{label} wrote to {table}"


def test_extraction_short_circuit_does_not_even_build_an_ai_client(client, db, monkeypatch):
    """Stronger than 'never called': the client is never constructed."""
    patch_session(monkeypatch, db)
    monkeypatch.setattr(
        api, "build_client", lambda: pytest.fail("build_client must not be reached")
    )

    response = upload(client, content=b"", filename="empty.pdf")

    assert response.status_code == 422


# ── 3. model says not_a_resume -> no writes ─────────────────────────────────


@pytest.mark.parametrize("model_status", ["not_a_resume", "unparseable"])
def test_model_rejecting_the_document_writes_nothing(model_status, client, db, monkeypatch):
    patch_session(monkeypatch, db)
    fake = FakeAI(json.dumps({"status": model_status, "profile": {}, "certifications": [],
                              "work_experience": [], "projects": []}))
    monkeypatch.setattr(api, "build_client", lambda: fake)

    response = upload(client)

    assert response.status_code == 200
    body = response.json()
    assert body["status"] == model_status
    assert body["written"] is None
    # The model WAS consulted -- this is a different failure from extraction.
    assert len(fake.calls) == 1
    for table in ("career_profiles", "certifications", "work_experience", "projects"):
        assert db.tables[table] == [], table


def test_non_ok_status_cannot_smuggle_rows_through(client, db, monkeypatch):
    """A not_a_resume payload carrying items must still write nothing."""
    patch_session(monkeypatch, db)
    payload = dict(GOOD_PARSE, status="not_a_resume")
    monkeypatch.setattr(api, "build_client", lambda: FakeAI(json.dumps(payload)))

    response = upload(client)

    assert response.json()["status"] == "not_a_resume"
    assert db.tables["certifications"] == []
    assert db.tables["career_profiles"] == []


# ── 4. malformed model output -> structured failure, never a 500 ────────────


@pytest.mark.parametrize(
    ("label", "text"),
    [
        ("not json", "I'm sorry, I can't help with that."),
        ("truncated json", '{"status": "ok", "certifications": ['),
        ("json array", "[1, 2, 3]"),
        ("empty string", ""),
        ("missing status", json.dumps({"profile": {}, "certifications": []})),
        ("unknown status", json.dumps({"status": "maybe", "profile": {}})),
        ("status not a string", json.dumps({"status": 7})),
    ],
)
def test_malformed_model_output_returns_structured_failure(label, text, client, db, monkeypatch):
    patch_session(monkeypatch, db)
    monkeypatch.setattr(api, "build_client", lambda: FakeAI(text))

    response = upload(client)

    assert response.status_code == 200, f"{label}: expected a structured failure, not an error code"
    assert response.status_code != 500
    body = response.json()
    assert body["status"] == "parse_failed", label
    assert body["written"] is None
    assert body["errors"] and isinstance(body["errors"][0], str)
    for table in ("career_profiles", "certifications", "work_experience", "projects"):
        assert db.tables[table] == [], f"{label} wrote to {table}"


def test_ai_request_error_is_also_a_structured_failure(client, db, monkeypatch):
    from GradusIQ_career.ai.errors import AIRequestError

    patch_session(monkeypatch, db)

    class Failing:
        def complete(self, **kwargs):
            raise AIRequestError("OpenRouter request failed: timeout")

    monkeypatch.setattr(api, "build_client", lambda: Failing())

    response = upload(client)

    assert response.status_code == 200
    assert response.json()["status"] == "parse_failed"
    assert "timeout" in response.json()["errors"][0]


def test_garbage_items_are_dropped_but_the_upload_still_succeeds(client, db, monkeypatch):
    """One bad entry must not discard a good resume."""
    patch_session(monkeypatch, db)
    payload = {
        "status": "ok",
        "profile": {"target_roles": ["SWE"], "interests": None,
                    "skills_technical": ["Python", "Python"], "skills_soft": "teamwork"},
        "certifications": [
            {"name": "Good Cert", "status": "expired"},   # status outside the CHECK
            {"issuer": "No Name Inc"},                     # missing required name
            "not an object",
        ],
        "work_experience": [{"employer": "Acme"}],
        "projects": "not a list",
    }
    monkeypatch.setattr(api, "build_client", lambda: FakeAI(json.dumps(payload)))

    response = upload(client)
    body = response.json()

    assert body["status"] == "ok"
    assert body["written"]["certifications"]["inserted"] == 1
    assert body["written"]["projects"]["inserted"] == 0
    cert = db.rows_for("certifications", STUDENT_A)[0]
    assert cert["name"] == "Good Cert"
    # Coerced to null rather than sent to Postgres to fail the CHECK.
    assert cert["status"] is None
    profile = db.tables["career_profiles"][0]
    assert profile["skills_technical"] == ["Python"]   # de-duplicated
    assert profile["skills_soft"] == ["teamwork"]      # bare string tolerated
    assert profile["interests"] == []
    assert body["warnings"], "dropped/coerced fields must be reported"


# ── 5. career_profiles bootstrap ────────────────────────────────────────────


def test_second_upload_does_not_modify_an_existing_career_profile(client, db, monkeypatch):
    patch_session(monkeypatch, db)
    monkeypatch.setattr(api, "build_client", lambda: FakeAI())

    first = upload(client)
    assert first.json()["career_profile"]["outcome"] == "created"
    original = dict(db.tables["career_profiles"][0])

    # A second resume whose profile differs in every scalar field.
    different = dict(
        GOOD_PARSE,
        profile={
            "target_roles": ["Product Manager"],
            "interests": ["design"],
            "skills_technical": ["Figma"],
            "skills_soft": ["leadership"],
        },
    )
    monkeypatch.setattr(api, "build_client", lambda: FakeAI(json.dumps(different)))

    second = upload(client)

    assert second.status_code == 200
    assert second.json()["career_profile"]["outcome"] == "already_existed_untouched"
    assert len(db.tables["career_profiles"]) == 1
    assert db.tables["career_profiles"][0] == original, (
        "an existing career_profiles row must never be rewritten by a later upload"
    )


def test_existing_confirmed_profile_is_not_demoted_by_a_new_upload(client, db, monkeypatch):
    """The row may already be student-confirmed; the parser must not touch it."""
    db.tables["career_profiles"].append(
        {
            "id": "cp-existing",
            "student_id": STUDENT_A,
            "target_roles": ["Data Analyst"],
            "interests": [],
            "skills_technical": [],
            "skills_soft": [],
            "source": "manual",
            "confirmed_at": "2026-01-01T00:00:00+00:00",
        }
    )
    patch_session(monkeypatch, db)
    monkeypatch.setattr(api, "build_client", lambda: FakeAI())

    response = upload(client)

    assert response.json()["career_profile"]["outcome"] == "already_existed_untouched"
    row = db.tables["career_profiles"][0]
    assert row["source"] == "manual"
    assert row["confirmed_at"] == "2026-01-01T00:00:00+00:00"
    assert row["target_roles"] == ["Data Analyst"]
    # Children still attach to it.
    assert db.rows_for("certifications", STUDENT_A)[0]["career_profile_id"] == "cp-existing"


# ── 6. duplicate child rows are skipped, not duplicated or overwritten ──────


def test_duplicate_rows_are_skipped_and_reported_separately(client, db, monkeypatch):
    patch_session(monkeypatch, db)
    monkeypatch.setattr(api, "build_client", lambda: FakeAI())

    first = upload(client)
    assert first.json()["written"]["certifications"] == {"inserted": 1, "skipped_duplicate": 0}
    before = {t: [dict(r) for r in db.rows_for(t, STUDENT_A)] for t in NATURAL_KEYS}

    second = upload(client)

    assert second.status_code == 200
    body = second.json()
    for table in NATURAL_KEYS:
        assert body["written"][table] == {"inserted": 0, "skipped_duplicate": 1}, table
        assert len(db.rows_for(table, STUDENT_A)) == 1, f"{table} duplicated"
        assert db.rows_for(table, STUDENT_A) == before[table], f"{table} row was modified"


def test_duplicate_detection_treats_null_role_as_equal(client, db, monkeypatch):
    """work_experience's index is NULLS NOT DISTINCT -- verified live in Stage 1."""
    patch_session(monkeypatch, db)
    payload = dict(GOOD_PARSE, work_experience=[{"employer": "Acme Corp", "role": None}])
    monkeypatch.setattr(api, "build_client", lambda: FakeAI(json.dumps(payload)))

    upload(client)
    assert len(db.rows_for("work_experience", STUDENT_A)) == 1

    second = upload(client)

    assert second.json()["written"]["work_experience"] == {"inserted": 0, "skipped_duplicate": 1}
    assert len(db.rows_for("work_experience", STUDENT_A)) == 1


def test_a_genuinely_new_item_still_inserts_alongside_a_duplicate(client, db, monkeypatch):
    patch_session(monkeypatch, db)
    monkeypatch.setattr(api, "build_client", lambda: FakeAI())
    upload(client)

    payload = dict(
        GOOD_PARSE,
        certifications=[
            {"name": "AWS Cloud Practitioner"},          # duplicate
            {"name": "Azure Fundamentals"},              # new
        ],
    )
    monkeypatch.setattr(api, "build_client", lambda: FakeAI(json.dumps(payload)))

    second = upload(client)

    assert second.json()["written"]["certifications"] == {"inserted": 1, "skipped_duplicate": 1}
    names = sorted(r["name"] for r in db.rows_for("certifications", STUDENT_A))
    assert names == ["AWS Cloud Practitioner", "Azure Fundamentals"]


# ── 7. confirm endpoint, including cross-student isolation ──────────────────


def seed_unconfirmed(db, student_id, profile_id):
    db.tables["career_profiles"].append(
        {"id": profile_id, "student_id": student_id, "source": "resume_parse",
         "confirmed_at": None, "target_roles": [], "interests": [],
         "skills_technical": [], "skills_soft": []}
    )
    for table, extra in (
        ("certifications", {"name": f"Cert {student_id[:4]}"}),
        ("work_experience", {"employer": f"Emp {student_id[:4]}", "role": None}),
        ("projects", {"name": f"Proj {student_id[:4]}"}),
    ):
        db.tables[table].append(
            {"id": f"{table}-{student_id[:4]}", "student_id": student_id,
             "career_profile_id": profile_id, "source": "resume_parse",
             "confirmed_at": None, **extra}
        )


def test_confirm_stamps_every_unconfirmed_row_for_the_caller_only(client, db, monkeypatch):
    seed_unconfirmed(db, STUDENT_A, "cp-a")
    seed_unconfirmed(db, STUDENT_B, "cp-b")
    patch_session(monkeypatch, db, student_id=STUDENT_A)

    response = client.post(CONFIRM, headers=HEADERS)

    assert response.status_code == 200, response.text
    body = response.json()
    assert body["status"] == "ok"
    assert body["scope"] == "all_unconfirmed"
    assert body["confirmed"] == {
        "career_profiles": 1, "certifications": 1, "work_experience": 1, "projects": 1
    }
    assert body["total_confirmed"] == 4

    # Student A: everything confirmed.
    for table in ("career_profiles", "certifications", "work_experience", "projects"):
        rows = db.rows_for(table, STUDENT_A)
        assert all(r["confirmed_at"] is not None for r in rows), table
        assert all(r.get("updated_at") for r in rows), f"{table} updated_at not maintained"

    # Student B: untouched. This is the isolation assertion.
    for table in ("career_profiles", "certifications", "work_experience", "projects"):
        rows = db.rows_for(table, STUDENT_B)
        assert rows, table
        assert all(r["confirmed_at"] is None for r in rows), (
            f"{table}: another student's rows were confirmed -- RLS/scoping leak"
        )


def test_confirm_does_not_rewrite_an_earlier_confirmation(client, db, monkeypatch):
    seed_unconfirmed(db, STUDENT_A, "cp-a")
    db.tables["certifications"][0]["confirmed_at"] = "2020-01-01T00:00:00+00:00"
    patch_session(monkeypatch, db, student_id=STUDENT_A)

    body = client.post(CONFIRM, headers=HEADERS).json()

    assert body["confirmed"]["certifications"] == 0
    assert db.tables["certifications"][0]["confirmed_at"] == "2020-01-01T00:00:00+00:00"


def test_confirm_accepts_a_subset_by_id(client, db, monkeypatch):
    seed_unconfirmed(db, STUDENT_A, "cp-a")
    patch_session(monkeypatch, db, student_id=STUDENT_A)

    response = client.post(
        CONFIRM, headers=HEADERS, json={"certifications": ["certifications-8f14"]}
    )

    body = response.json()
    assert body["scope"] == "selection"
    assert body["confirmed"] == {
        "career_profiles": 0, "certifications": 1, "work_experience": 0, "projects": 0
    }
    assert db.rows_for("certifications", STUDENT_A)[0]["confirmed_at"] is not None
    assert db.rows_for("projects", STUDENT_A)[0]["confirmed_at"] is None
    assert db.tables["career_profiles"][0]["confirmed_at"] is None


def test_confirm_subset_cannot_reach_another_students_row_by_id(client, db, monkeypatch):
    """Naming B's id from A's session confirms nothing."""
    seed_unconfirmed(db, STUDENT_A, "cp-a")
    seed_unconfirmed(db, STUDENT_B, "cp-b")
    patch_session(monkeypatch, db, student_id=STUDENT_A)

    body = client.post(
        CONFIRM, headers=HEADERS, json={"certifications": ["certifications-1a2b"]}
    ).json()

    assert body["confirmed"]["certifications"] == 0
    assert db.rows_for("certifications", STUDENT_B)[0]["confirmed_at"] is None


def test_confirm_with_empty_body_confirms_everything(client, db, monkeypatch):
    seed_unconfirmed(db, STUDENT_A, "cp-a")
    patch_session(monkeypatch, db, student_id=STUDENT_A)

    body = client.post(CONFIRM, headers=HEADERS, json={}).json()

    assert body["scope"] == "all_unconfirmed"
    assert body["total_confirmed"] == 4


def test_confirm_writes_reviewed_academic_facts_without_touching_declarations(
    client, db, monkeypatch
):
    student = db.tables["students"][0]
    student.update(
        {
            "major_current": "Electrical Engineering",
            "major_intended": "N/A",
            "expected_graduation": "Fall 2030",
        }
    )
    seed_unconfirmed(db, STUDENT_A, "cp-a")
    profile = db.tables["career_profiles"][0]
    profile.update({"target_roles": [], "interests": [], "ai_anxiety_level": None})
    patch_session(monkeypatch, db, student_id=STUDENT_A)

    response = client.post(
        CONFIRM,
        headers=HEADERS,
        json={
            "academics": {
                "major_current": "Computer Engineering",
                "expected_graduation": "Spring 2029",
            }
        },
    )

    assert response.status_code == 200, response.text
    assert response.json()["academic_rows_updated"] == 1
    assert student["major_current"] == "Computer Engineering"
    assert student["expected_graduation"] == "Spring 2029"
    assert student["major_intended"] == "N/A"
    assert profile["target_roles"] == []
    assert profile["interests"] == []
    assert profile["ai_anxiety_level"] is None


def test_confirm_omitted_academic_facts_do_not_clear_existing_values(client, db, monkeypatch):
    student = db.tables["students"][0]
    student.update(
        {
            "major_current": "Computer Science",
            "major_intended": "Mathematics",
            "expected_graduation": "Fall 2028",
        }
    )
    seed_unconfirmed(db, STUDENT_A, "cp-a")
    patch_session(monkeypatch, db, student_id=STUDENT_A)

    response = client.post(CONFIRM, headers=HEADERS, json={"academics": {}})

    assert response.status_code == 200
    assert response.json()["academic_rows_updated"] == 0
    assert student["major_current"] == "Computer Science"
    assert student["major_intended"] == "Mathematics"
    assert student["expected_graduation"] == "Fall 2028"


def test_resume_confirmation_flows_into_canonical_dashboard_fields(client, db, monkeypatch):
    patch_session(monkeypatch, db, student_id=STUDENT_A)
    monkeypatch.setattr(api, "build_client", lambda: FakeAI())

    uploaded = upload(
        client,
        content=make_pdf(
            [
                "Bachelor of Science in Computer Engineering",
                "Expected Graduation: May 2029",
                "Skills: Python, TypeScript",
                "Praxigen - Software Engineering Intern",
            ]
        ),
    )
    facts = uploaded.json()["academics"]
    confirmed = client.post(CONFIRM, headers=HEADERS, json={"academics": facts})

    assert confirmed.status_code == 200, confirmed.text
    canonical = build_student_intelligence_profile(FakeSupabase(db, STUDENT_A), STUDENT_A)
    assert canonical.academics.summary.major_current == "Computer Engineering"
    assert canonical.identity.expected_graduation == "Spring 2029"
    assert canonical.career.confirmed is True
    assert canonical.career.skills.technical == ["Python", "SQL"]
    display_major = (
        canonical.academics.summary.major_current
        or canonical.academics.summary.major_intended
        or "Major not provided"
    )
    assert display_major == "Computer Engineering"


def test_confirm_requires_a_session(client, db, monkeypatch):
    monkeypatch.setattr(
        api, "build_client_for_token", lambda t: pytest.fail("must not build a client")
    )
    assert client.post(CONFIRM, headers=PROXY_HEADERS).status_code == 401


def test_confirm_404s_when_no_student_row_is_visible(client, db, monkeypatch):
    empty = FakeDB()  # no students at all
    monkeypatch.setattr(api, "build_client_for_token", lambda t: FakeSupabase(empty, STUDENT_A))

    assert client.post(CONFIRM, headers=HEADERS).status_code == 404


# ── 8. the proxy-secret dependency (and so the rate limit) is really wired ──


def test_upload_requires_the_proxy_secret(db, monkeypatch):
    unauth = TestClient(api.create_app(make_test_config()))
    patch_session(monkeypatch, db)
    monkeypatch.setattr(api, "build_client", lambda: ExplodingAI())

    response = unauth.post(
        UPLOAD, headers=AUTH, files={"file": ("r.pdf", make_pdf(), PDF)}
    )

    assert response.status_code == 401


def test_confirm_requires_the_proxy_secret(db, monkeypatch):
    unauth = TestClient(api.create_app(make_test_config()))
    assert unauth.post(CONFIRM, headers=AUTH).status_code == 401


@pytest.mark.parametrize("path", [UPLOAD, CONFIRM])
def test_rate_limit_actually_applies_to_these_routes(path, db, monkeypatch):
    """Exhausting the shared limiter must 429 -- proof the dependency runs."""
    limited = TestClient(
        api.create_app(make_test_config(rate_limit_requests=1)), headers=PROXY_HEADERS
    )
    seed_unconfirmed(db, STUDENT_A, "cp-a")
    patch_session(monkeypatch, db)
    monkeypatch.setattr(api, "build_client", lambda: FakeAI())

    def call():
        if path == UPLOAD:
            return limited.post(UPLOAD, headers=HEADERS,
                                files={"file": ("r.pdf", make_pdf(), PDF)})
        return limited.post(CONFIRM, headers=HEADERS)

    first = call()
    second = call()

    assert first.status_code != 429, first.text
    assert second.status_code == 429
    assert second.json()["detail"] == "Request rate limit exceeded."


def test_upload_is_inside_the_ai_concurrency_gate(db, monkeypatch):
    """Capacity 0 -> 429 before the model is reached."""
    gated = TestClient(
        api.create_app(make_test_config(max_concurrent_ai_requests=0)), headers=PROXY_HEADERS
    )
    patch_session(monkeypatch, db)
    fake = FakeAI()
    monkeypatch.setattr(api, "build_client", lambda: fake)

    response = gated.post(UPLOAD, headers=HEADERS, files={"file": ("r.pdf", make_pdf(), PDF)})

    assert response.status_code == 429
    assert response.json()["detail"] == "AI service is busy; retry later."
    assert fake.calls == [], "the gate must wrap the model call, not follow it"
    assert db.tables["career_profiles"] == []


def test_upload_route_declares_the_proxy_dependency():
    """The wiring itself, asserted directly alongside the behavioral tests."""
    routes = {
        r.path: r for r in api.router.routes if getattr(r, "path", None) in (UPLOAD, CONFIRM)
    }
    assert set(routes) == {UPLOAD, CONFIRM}
    for path, route in routes.items():
        names = [getattr(d.dependency, "__name__", "") for d in route.dependencies]
        assert "authorize_proxy_request" in names, path
        assert "POST" in route.methods, path


# ── 9. SUPABASE_SECRET_KEY is never read on either path ─────────────────────


@pytest.mark.parametrize("path", [UPLOAD, CONFIRM])
def test_supabase_secret_key_is_never_read(path, db, monkeypatch):
    """Non-vacuous: build_client_for_token runs for real under the spy.

    Only the supabase SDK's create_client is replaced, so
    build_client_for_token's own _required_env calls execute and are visible
    to the spy -- the same shape as tests/test_api.py:809. Stubbing
    build_client_for_token itself (as the lighter tests above do) would make
    this assertion vacuous.
    """
    import os

    from GradusIQ_career import supabase_client as supabase_client_module

    seed_unconfirmed(db, STUDENT_A, "cp-a")

    reads: list[str] = []
    original_get = os.environ.get

    def spying_get(key, *args, **kwargs):
        reads.append(key)
        if key == "SUPABASE_SECRET_KEY":
            raise AssertionError("SUPABASE_SECRET_KEY must never be read by these routes")
        return original_get(key, *args, **kwargs)

    monkeypatch.setenv("SUPABASE_URL", "https://example.supabase.co")
    monkeypatch.setenv("SUPABASE_PUBLISHABLE_KEY", "test-publishable-key")
    monkeypatch.setattr(
        supabase_client_module, "create_client", lambda url, key: FakeSupabase(db, STUDENT_A)
    )
    monkeypatch.setattr(api, "build_client", lambda: FakeAI())
    monkeypatch.setattr(os.environ, "get", spying_get)

    live = TestClient(api.create_app(make_test_config()), headers=PROXY_HEADERS)
    if path == UPLOAD:
        response = live.post(UPLOAD, headers=HEADERS, files={"file": ("r.pdf", make_pdf(), PDF)})
    else:
        response = live.post(CONFIRM, headers=HEADERS)

    assert response.status_code == 200, response.text
    # Non-vacuity: the real builder ran and read the two keys it is allowed to.
    assert "SUPABASE_URL" in reads
    assert "SUPABASE_PUBLISHABLE_KEY" in reads
    assert "SUPABASE_SECRET_KEY" not in reads


@pytest.mark.parametrize("path", [UPLOAD, CONFIRM])
def test_caller_token_is_handed_to_postgrest(path, db, monkeypatch):
    """RLS only applies if the token actually reaches postgrest.auth()."""
    from GradusIQ_career import supabase_client as supabase_client_module

    seed_unconfirmed(db, STUDENT_A, "cp-a")
    built: list[FakeSupabase] = []

    def fake_create_client(url, key):
        instance = FakeSupabase(db, STUDENT_A)
        built.append(instance)
        return instance

    monkeypatch.setenv("SUPABASE_URL", "https://example.supabase.co")
    monkeypatch.setenv("SUPABASE_PUBLISHABLE_KEY", "test-publishable-key")
    monkeypatch.setattr(supabase_client_module, "create_client", fake_create_client)
    monkeypatch.setattr(api, "build_client", lambda: FakeAI())

    live = TestClient(api.create_app(make_test_config()), headers=PROXY_HEADERS)
    if path == UPLOAD:
        live.post(UPLOAD, headers=HEADERS, files={"file": ("r.pdf", make_pdf(), PDF)})
    else:
        live.post(CONFIRM, headers=HEADERS)

    assert built, "no client was built"
    assert built[0].postgrest.tokens == ["real-session-jwt"]


# ── shared session-auth behavior ────────────────────────────────────────────


@pytest.mark.parametrize("path", [UPLOAD, CONFIRM])
def test_no_authorization_header_is_401(path, client, monkeypatch):
    monkeypatch.setattr(
        api, "build_client_for_token", lambda t: pytest.fail("must not build a client")
    )
    kwargs = {"headers": PROXY_HEADERS}
    if path == UPLOAD:
        kwargs["files"] = {"file": ("r.pdf", make_pdf(), PDF)}
    assert client.post(path, **kwargs).status_code == 401


@pytest.mark.parametrize("path", [UPLOAD, CONFIRM])
def test_supabase_config_error_is_503(path, client, monkeypatch):
    from GradusIQ_career.supabase_client import SupabaseConfigError

    monkeypatch.setattr(
        api,
        "build_client_for_token",
        lambda t: (_ for _ in ()).throw(SupabaseConfigError("SUPABASE_URL is not set.")),
    )
    kwargs = {"headers": HEADERS}
    if path == UPLOAD:
        kwargs["files"] = {"file": ("r.pdf", make_pdf(), PDF)}

    response = client.post(path, **kwargs)

    assert response.status_code == 503
    assert "SUPABASE_URL" in response.json()["detail"]


def test_upload_404s_when_no_student_row_is_visible(client, monkeypatch):
    empty = FakeDB()
    monkeypatch.setattr(api, "build_client_for_token", lambda t: FakeSupabase(empty, STUDENT_A))
    monkeypatch.setattr(api, "build_client", lambda: ExplodingAI())

    response = upload(client)

    assert response.status_code == 404


def test_oversized_upload_is_413(client, db, monkeypatch):
    patch_session(monkeypatch, db)
    monkeypatch.setattr(api, "build_client", lambda: ExplodingAI())

    oversized = b"%PDF-1.7\n" + b"\x00" * (api.MAX_RESUME_BYTES + 10)
    response = upload(client, content=oversized)

    assert response.status_code == 413
    assert db.tables["career_profiles"] == []
